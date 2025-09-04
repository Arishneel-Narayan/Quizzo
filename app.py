import streamlit as st
import pandas as pd
import markdown2
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

def add_table_to_docx(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    return doc

def markdown_to_docx(doc, md_text):
    html = markdown2.markdown(md_text)
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    for elem in soup.children:
        if elem.name == 'h1':
            p = doc.add_paragraph(elem.text)
            p.style = doc.styles['Heading 1']
        elif elem.name == 'h2':
            p = doc.add_paragraph(elem.text)
            p.style = doc.styles['Heading 2']
        elif elem.name == 'h3':
            p = doc.add_paragraph(elem.text)
            p.style = doc.styles['Heading 3']
        elif elem.name == 'ul':
            for li in elem.find_all('li'):
                p = doc.add_paragraph(li.text, style='List Bullet')
        elif elem.name == 'ol':
            for li in elem.find_all('li'):
                p = doc.add_paragraph(li.text, style='List Number')
        elif elem.name == 'table':
            # skip, handled by CSV
            continue
        elif elem.name == 'p':
            doc.add_paragraph(elem.text)
    return doc

def generate_docx(md_file, csv_file):
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Parse markdown
    md_text = md_file.read().decode('utf-8')
    doc = markdown_to_docx(doc, md_text)
    # Add CSV as table
    if csv_file is not None:
        df = pd.read_csv(csv_file)
        doc.add_paragraph('')
        doc.add_paragraph('Data Table:', style='Heading 2')
        doc = add_table_to_docx(doc, df)
    return doc

def docx_to_pdf(docx_bytes):
    if not DOCX2PDF_AVAILABLE:
        return None
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
        tmp_docx.write(docx_bytes.getvalue())
        tmp_docx.flush()
        pdf_path = tmp_docx.name.replace('.docx', '.pdf')
        docx2pdf_convert(tmp_docx.name, pdf_path)
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        os.remove(tmp_docx.name)
        os.remove(pdf_path)
    return pdf_bytes

def main():
    st.title('Professional Markdown & CSV to Word/PDF Exporter')
    st.write('Upload a Markdown file and a CSV file. Choose your output format for a high-quality, professional document.')
    md_file = st.file_uploader('Upload Markdown file', type=['md', 'markdown'])
    csv_file = st.file_uploader('Upload CSV file (optional)', type=['csv'])
    output_format = st.selectbox('Select output format', ['Word (.docx)', 'PDF (.pdf)'])
    if st.button('Generate Document'):
        if md_file is None:
            st.error('Please upload a Markdown file.')
            return
        doc = generate_docx(md_file, csv_file)
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        if output_format == 'Word (.docx)':
            st.download_button('Download Word Document', docx_bytes, file_name='output.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            if not DOCX2PDF_AVAILABLE:
                st.warning('PDF export requires docx2pdf. Please install it with: pip install docx2pdf')
            else:
                pdf_bytes = docx_to_pdf(docx_bytes)
                if pdf_bytes:
                    st.download_button('Download PDF Document', pdf_bytes, file_name='output.pdf', mime='application/pdf')
                else:
                    st.error('PDF conversion failed.')

if __name__ == '__main__':
    main()
